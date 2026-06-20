# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, os

with open(r"C:\Users\周进法\Desktop\ai-novel\output\chapters.json", "r", encoding="utf-8") as f:
    chapters = json.load(f)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

title = doc.add_heading('\u96f7\u52ab\u6dec\u4f53', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('AI Novel Studio \u751f\u6210 \u00b7 DeepSeek \u9a71\u52a8')
run.font.size = Pt(12)

doc.add_paragraph()

for ch in chapters:
    doc.add_heading(ch['title'], level=1)
    for para_text in ch['content'].split('\n\n'):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

doc.save(r"C:\Users\周进法\Desktop\ai-novel\output\雷劫淬体_第1-5章.docx")
print('DOCX saved OK')