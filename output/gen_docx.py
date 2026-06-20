# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

with open(r"C:\Users\周进法\Desktop\ai-novel\output\chapters.json", "r", encoding="utf-8") as f:
    chapters = json.load(f)

doc = Document()
doc.styles["Normal"].font.name = "Microsoft YaHei"
doc.styles["Normal"].font.size = Pt(11)

t = doc.add_heading("雷劫淬体", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("AI Novel Studio 生成")
r.font.size = Pt(11)

doc.add_paragraph()

for ch in chapters:
    doc.add_heading(ch["title"], 1)
    for pt in ch["content"].split("\n\n"):
        pt = pt.strip()
        if pt:
            p = doc.add_paragraph(pt)
            for r in p.runs:
                r.font.name = "Microsoft YaHei"
                r.font.size = Pt(11)

path = r"C:\Users\周进法\Desktop\ai-novel\output\雷劫淬体_第1-5章.docx"
doc.save(path)
print("DOCX saved:", path)